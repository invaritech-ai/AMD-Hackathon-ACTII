"""Extract clean text from *text-native* documents, preserving tables.

    text-layer PDF        -> pdfplumber (text + text-strategy tables)
    .docx                 -> python-docx (paragraphs + tables)
    .xlsx                 -> openpyxl (rows)
    .csv                  -> stdlib csv (rows)
    .txt / .md / .json    -> decode

Images and scanned PDFs (no usable text layer) carry no native text — those go
to the vision-LLM OCR path (see ``services.vision_ocr``), gated by a cheap
classification pass so junk never burns a transcription. ``extract_native_text``
returns ``None`` for them to signal "needs OCR".

Tables are the point: financial docs live and die on their line-item tables, so
every tabular source is rendered to Markdown. Every extractor is
dependency-guarded and never raises — a failed extract returns "" (or None).
"""

from __future__ import annotations

import csv
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# A PDF text layer shorter than this is treated as scanned -> route to OCR.
_MIN_TEXT_LAYER_CHARS = 50

_IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"}
_TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".json", ".log", ".xml", ".yaml", ".yml"}

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".xlsx", ".csv", *_IMAGE_SUFFIXES, *_TEXT_SUFFIXES}


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_SUFFIXES


def extract_native_text(file_path: Path) -> str | None:
    """Clean text for text-native formats. Never raises.

    Returns ``None`` when the document has no native text layer (image, or a
    scanned PDF) — the caller routes those to vision OCR.
    """
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".pdf":
            text = _pdf_text_layer(file_path)
            return text if len(text.strip()) >= _MIN_TEXT_LAYER_CHARS else None
        if suffix in _IMAGE_SUFFIXES:
            return None
        if suffix == ".docx":
            return _from_docx(file_path)
        if suffix == ".xlsx":
            return _from_xlsx(file_path)
        if suffix == ".csv":
            return _from_csv(file_path)
        if suffix in _TEXT_SUFFIXES:
            return _decode(file_path)
        return _decode(file_path)
    except Exception:  # never let extraction 500 the request
        logger.exception("native text extraction failed for %s", file_path.name)
        return None


# ── Markdown helpers ──────────────────────────────────────────────────

def rows_to_markdown(rows: list[list[str]]) -> str:
    """Render rows as a GitHub Markdown table. First row is the header."""
    cleaned = [
        [(c if c is not None else "").replace("\n", " ").replace("|", "\\|").strip() for c in row]
        for row in rows
    ]
    cleaned = [r for r in cleaned if any(cell for cell in r)]
    if not cleaned:
        return ""
    width = max(len(r) for r in cleaned)
    cleaned = [r + [""] * (width - len(r)) for r in cleaned]
    header, *body = cleaned
    lines = ["| " + " | ".join(header) + " |", "|" + "|".join(["---"] * width) + "|"]
    lines += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(lines)


# ── PDF (text layer only) ─────────────────────────────────────────────

def _pdf_text_layer(file_path: Path) -> str:
    import pdfplumber

    text_parts: list[str] = []
    table_strategy = {"vertical_strategy": "text", "horizontal_strategy": "text"}
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
            for table in page.extract_tables(table_strategy):
                md = rows_to_markdown(table)
                if md:
                    text_parts.append(md)
    return "\n\n".join(p for p in text_parts if p.strip())


# ── Office formats ────────────────────────────────────────────────────

def _from_docx(file_path: Path) -> str:
    import docx

    document = docx.Document(str(file_path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        md = rows_to_markdown(rows)
        if md:
            parts.append(md)
    return "\n\n".join(parts)


def _from_xlsx(file_path: Path) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        rows = [
            [("" if v is None else str(v)) for v in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        md = rows_to_markdown(rows)
        if md:
            parts.append(f"## {sheet.title}\n\n{md}")
    workbook.close()
    return "\n\n".join(parts)


def _from_csv(file_path: Path) -> str:
    with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
        rows = list(csv.reader(f))
    return rows_to_markdown(rows)


# ── Plain text ────────────────────────────────────────────────────────

def _decode(file_path: Path) -> str:
    data = file_path.read_bytes()
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        text = data.decode("utf-8", errors="replace")
    # Pretty-print JSON so keys are on their own lines (helps keyword classification).
    if file_path.suffix.lower() == ".json":
        try:
            return json.dumps(json.loads(text), indent=2)
        except json.JSONDecodeError:
            pass
    return text
